# -*- coding: utf-8 -*-
"""
csv_para_word.py
────────────────
Converte automaticamente todos os arquivos CSV gerados pelo pipeline
(em resultados/tabelas/) para tabelas formatadas em Word (.docx).

Modos de uso:
  1. Um único documento com todas as tabelas (padrão):
       python csv_para_word.py

  2. Um documento Word por CSV:
       python csv_para_word.py --separados

  3. Converter apenas um CSV específico:
       python csv_para_word.py --arquivo resultados/tabelas/top_20_apps.csv
"""

import argparse
import sys
from pathlib import Path
from datetime import datetime

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── caminhos ──────────────────────────────────────────────────────────────────
ROOT_DIR    = Path(__file__).resolve().parent
TABELAS_DIR = ROOT_DIR / "resultados" / "tabelas"
SAIDA_DIR   = ROOT_DIR / "resultados" / "tabelas_word"
SAIDA_DIR.mkdir(parents=True, exist_ok=True)

# ── mapeamento: nome do arquivo → título legível ───────────────────────────
TITULOS = {
    "top_20_apps.csv":                "Top 20 Aplicativos por Instalações",
    "estatistica_descritiva.csv":     "Estatística Descritiva das Variáveis Numéricas",
    "correlacoes.csv":                "Análise de Correlação entre Variáveis",
    "distribuicao_desenvolvedor.csv": "Distribuição por Tipo de Desenvolvedor",
    "sentimentos.csv":                "Análise de Sentimentos das Avaliações",
    "distribuicao_tematica.csv":      "Distribuição Temática das Avaliações",
    "topicos_lda.csv":                "Modelagem de Tópicos (LDA)",
    "frequencia_palavras.csv":        "Frequência de Palavras nas Avaliações",
}

# Ordem de exibição preferencial (outros CSVs são adicionados ao final)
ORDEM = [
    "top_20_apps.csv",
    "estatistica_descritiva.csv",
    "distribuicao_desenvolvedor.csv",
    "correlacoes.csv",
    "sentimentos.csv",
    "distribuicao_tematica.csv",
    "topicos_lda.csv",
    "frequencia_palavras.csv",
]


# ════════════════════════════════════════════════════════════════════════════
# UTILITÁRIOS DE FORMATAÇÃO
# ════════════════════════════════════════════════════════════════════════════

def _set_doc_margins(doc: Document, margin_cm: float = 2.5):
    """Define margens do documento."""
    for section in doc.sections:
        section.top_margin    = Cm(margin_cm)
        section.bottom_margin = Cm(margin_cm)
        section.left_margin   = Cm(margin_cm)
        section.right_margin  = Cm(margin_cm)


def _shade_cell(cell, fill: str):
    """Aplica cor de fundo (hex sem #) em uma célula."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _add_heading(doc: Document, text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)
    return h


def _add_paragraph(doc: Document, text: str, bold: bool = False, italic: bool = False,
                   align=WD_ALIGN_PARAGRAPH.LEFT, size: int = 11) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(6)


def _titulo_amigavel(nome_arquivo: str) -> str:
    """Retorna título legível para o CSV, ou converte o nome automaticamente."""
    if nome_arquivo in TITULOS:
        return TITULOS[nome_arquivo]
    # fallback: remove extensão, substitui _ por espaço, capitaliza
    return nome_arquivo.replace(".csv", "").replace("_", " ").title()


# ════════════════════════════════════════════════════════════════════════════
# INSERÇÃO DE TABELA CSV
# ════════════════════════════════════════════════════════════════════════════

def _inserir_tabela(doc: Document, df: pd.DataFrame, max_rows: int = 500) -> None:
    """Insere DataFrame como tabela Word formatada."""
    df = df.head(max_rows).reset_index(drop=True)

    if df.empty:
        _add_paragraph(doc, "(Tabela vazia — sem dados.)", italic=True)
        return

    num_cols = len(df.columns)
    table = doc.add_table(rows=1, cols=num_cols, style="Table Grid")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # ── cabeçalho ──────────────────────────────────────────────────────────
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        cell = hdr_cells[i]
        cell.text = str(col)
        _shade_cell(cell, "1A237E")          # azul escuro
        for par in cell.paragraphs:
            for run in par.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)  # branco
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── linhas de dados ────────────────────────────────────────────────────
    for idx, (_, row) in enumerate(df.iterrows()):
        cells = table.add_row().cells
        fill = "E8EAF6" if idx % 2 == 0 else "FFFFFF"   # zebra
        for i, val in enumerate(row):
            cell = cells[i]
            cell.text = "" if pd.isna(val) else str(val)
            _shade_cell(cell, fill)
            for par in cell.paragraphs:
                for run in par.runs:
                    run.font.size = Pt(8)

    # legenda de linhas truncadas
    total = len(df)
    if total == max_rows:
        doc.add_paragraph()
        _add_paragraph(doc,
            f"* Tabela limitada a {max_rows} linhas por legibilidade.",
            italic=True, size=9)


# ════════════════════════════════════════════════════════════════════════════
# CAPA E RODAPÉ
# ════════════════════════════════════════════════════════════════════════════

def _gerar_capa(doc: Document, titulo_doc: str) -> None:
    _set_doc_margins(doc)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for _ in range(5):
        doc.add_paragraph()

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run(titulo_doc)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Pipeline de Infodemiologia — Google Play Store")
    r.italic = True
    r.font.size = Pt(13)

    for _ in range(6):
        doc.add_paragraph()

    data = doc.add_paragraph()
    data.alignment = WD_ALIGN_PARAGRAPH.CENTER
    data.add_run(f"Gerado em: {datetime.now().strftime('%d/%m/%Y %H:%M')}").font.size = Pt(11)

    doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
# DOCUMENTO ÚNICO (todas as tabelas)
# ════════════════════════════════════════════════════════════════════════════

def gerar_documento_unico(csvs: list[Path]) -> Path:
    """Gera um único .docx com todas as tabelas CSV."""
    doc = Document()
    _gerar_capa(doc, "TABELAS DE RESULTADOS\nINFODEMIOLOGIA E INFOMETRIA")

    # sumário automático
    _add_heading(doc, "ÍNDICE DE TABELAS", 1)
    for i, csv in enumerate(csvs, 1):
        doc.add_paragraph(f"{i}. {_titulo_amigavel(csv.name)}")
    doc.add_page_break()

    for i, csv_path in enumerate(csvs, 1):
        titulo = _titulo_amigavel(csv_path.name)
        print(f"  [{i}/{len(csvs)}] {titulo} ← {csv_path.name}")

        _add_heading(doc, f"Tabela {i} — {titulo}", 1)
        _add_paragraph(doc, f"Arquivo de origem: {csv_path.name}", italic=True, size=9)
        doc.add_paragraph()

        try:
            df = pd.read_csv(csv_path, encoding="utf-8-sig")
            _add_paragraph(doc,
                f"{len(df)} linha(s) × {len(df.columns)} coluna(s)",
                size=9, italic=True)
            doc.add_paragraph()
            _inserir_tabela(doc, df)
        except Exception as exc:
            _add_paragraph(doc, f"[Erro ao ler arquivo: {exc}]", italic=True)

        doc.add_page_break()

    nome_saida = SAIDA_DIR / f"todas_tabelas_{datetime.now():%Y%m%d_%H%M%S}.docx"
    doc.save(str(nome_saida))
    return nome_saida


# ════════════════════════════════════════════════════════════════════════════
# DOCUMENTOS SEPARADOS (um por CSV)
# ════════════════════════════════════════════════════════════════════════════

def gerar_documentos_separados(csvs: list[Path]) -> list[Path]:
    """Gera um .docx independente para cada CSV."""
    saidas = []
    for csv_path in csvs:
        titulo = _titulo_amigavel(csv_path.name)
        print(f"  → {titulo}")

        doc = Document()
        _gerar_capa(doc, titulo)
        _set_doc_margins(doc)

        _add_heading(doc, titulo, 1)
        _add_paragraph(doc, f"Arquivo de origem: {csv_path.name}", italic=True, size=9)
        doc.add_paragraph()

        try:
            df = pd.read_csv(csv_path, encoding="utf-8-sig")
            _add_paragraph(doc,
                f"{len(df)} linha(s) × {len(df.columns)} coluna(s)",
                size=9, italic=True)
            doc.add_paragraph()
            _inserir_tabela(doc, df)
        except Exception as exc:
            _add_paragraph(doc, f"[Erro ao ler arquivo: {exc}]", italic=True)

        stem = csv_path.stem
        nome_saida = SAIDA_DIR / f"{stem}.docx"
        doc.save(str(nome_saida))
        saidas.append(nome_saida)

    return saidas


# ════════════════════════════════════════════════════════════════════════════
# PONTO DE ENTRADA
# ════════════════════════════════════════════════════════════════════════════

def _coletar_csvs(pasta: Path) -> list[Path]:
    """Retorna os CSVs na pasta, na ordem preferencial, depois os restantes."""
    todos = {p.name: p for p in sorted(pasta.glob("*.csv"))}
    ordenados = [todos.pop(nome) for nome in ORDEM if nome in todos]
    ordenados += sorted(todos.values(), key=lambda p: p.name)
    return ordenados


def main():
    parser = argparse.ArgumentParser(
        description="Converte CSVs do pipeline para Word automaticamente.")
    parser.add_argument("--separados", action="store_true",
                        help="Gera um Word por CSV (em vez de um único documento).")
    parser.add_argument("--arquivo", type=str, default=None,
                        help="Caminho para um único CSV a converter.")
    args = parser.parse_args()

    print("=" * 60)
    print("  CONVERSOR CSV → WORD")
    print("=" * 60)

    if args.arquivo:
        csv_path = Path(args.arquivo)
        if not csv_path.exists():
            print(f"ERRO: arquivo não encontrado: {csv_path}")
            sys.exit(1)
        csvs = [csv_path]
    else:
        csvs = _coletar_csvs(TABELAS_DIR)
        if not csvs:
            print(f"Nenhum CSV encontrado em: {TABELAS_DIR}")
            sys.exit(0)

    print(f"\nCSVs encontrados: {len(csvs)}")
    for c in csvs:
        print(f"  • {c.name}")
    print()

    if args.separados or args.arquivo:
        saidas = gerar_documentos_separados(csvs)
        print(f"\n✔ {len(saidas)} arquivo(s) Word gerado(s) em:")
        for s in saidas:
            print(f"  {s}")
    else:
        saida = gerar_documento_unico(csvs)
        print(f"\n✔ Documento único gerado:")
        print(f"  {saida}")

    print("\nConcluído.")


if __name__ == "__main__":
    main()
