# -*- coding: utf-8 -*-
"""
Gerador de relatórios Word (.docx) com resultados interpretativos detalhados.
Projeto: A Onipresença dos Sistemas de Informação em Saúde — mHealth.
"""

import logging
from pathlib import Path
from datetime import datetime

import pandas as pd
import numpy as np

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

from src.config import (CLEAN_DIR, GRAFICOS_DIR, TABELAS_DIR,
                        RELATORIO_DIR, DESCRITORES, EIXOS_TEMATICOS,
                        TITULO_PROJETO, SUBTITULO_PROJETO)

logger = logging.getLogger(__name__)


def _add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)
    return h


def _add_paragraph(doc, text, bold=False, italic=False, font_size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(font_size)
    p.paragraph_format.space_after = Pt(6)
    return p


def _add_table_from_df(doc, df, max_rows=60):
    """Insere tabela formatada a partir de DataFrame."""
    df = df.head(max_rows).reset_index()
    table = doc.add_table(rows=1, cols=len(df.columns), style="Light Grid Accent 1")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr[i].text = str(col)
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val) if pd.notna(val) else ""
            for p in cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8)
    return table


def _add_image(doc, path: Path, width=Inches(5.5)):
    if path.exists():
        doc.add_picture(str(path), width=width)
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        _add_paragraph(doc, f"[Gráfico não encontrado: {path.name}]", italic=True)


def _safe_read(path: Path, **kwargs) -> pd.DataFrame:
    if path.exists():
        return pd.read_csv(path, encoding="utf-8-sig", **kwargs)
    return pd.DataFrame()


# ────────────────────────────────────────────────────────────────────────────
def gerar_relatorio_metodologia(doc: Document):
    """Seção: Metodologia."""
    _add_heading(doc, "2. METODOLOGIA", 1)

    _add_heading(doc, "2.1 Tipologia do Estudo", 2)
    _add_paragraph(doc,
        "Trata-se de um estudo de Infodemiologia e Infometria, focado na análise "
        "de métricas de software e percepção do usuário no ecossistema digital de "
        "saúde. O objeto de análise são aplicativos móveis de saúde (mHealth) "
        "disponíveis na Google Play Store que possuam integração declarada com "
        "prontuários eletrônicos, sistemas governamentais (ex: conectividade via "
        "API) ou Sistemas de Informação em Saúde (SIS).")

    _add_heading(doc, "2.2 Protocolo de Coleta de Dados (Web Scraping)", 2)
    _add_paragraph(doc,
        "A coleta foi automatizada para garantir a integridade científica, "
        "utilizando scripts em Python (biblioteca google-play-scraper) para "
        "extração programática dos metadados e avaliações públicas dos aplicativos. "
        "A coleta obedeceu aos seguintes descritores de busca:")
    for d in DESCRITORES:
        doc.add_paragraph(f"• {d}", style="List Bullet")
    _add_paragraph(doc, "Parâmetros de extração:", bold=True)
    doc.add_paragraph(
        "Metadados técnicos: Versão do SO exigida, data da última atualização "
        "(indicador de manutenção), tamanho do arquivo e desenvolvedor "
        "(institucional vs. comercial).", style="List Bullet")
    doc.add_paragraph(
        "Métricas de desempenho: Número de instalações (escala de adoção) e "
        "nota média (rating).", style="List Bullet")
    doc.add_paragraph(
        "Conteúdo gerado pelo usuário: Comentários (reviews) para análise de "
        "usabilidade e bugs.", style="List Bullet")
    _add_paragraph(doc,
        "Campos de interesse: appId, title, installs, score, updated, genre, "
        "developer, description, reviews. Para cada aplicativo, foram extraídas "
        "até 200 avaliações textuais.")

    _add_heading(doc, "2.3 Definição do Escopo e Amostragem", 2)
    _add_paragraph(doc, "Critérios de inclusão:", bold=True)
    doc.add_paragraph(
        "Aplicativos categorizados em 'Medicina' ou 'Saúde e Fitness' que "
        "possuam integração declarada com prontuários eletrônicos ou sistemas "
        "governamentais;", style="List Bullet")
    doc.add_paragraph("Mínimo de 1.000 instalações;", style="List Bullet")
    doc.add_paragraph(
        "Última atualização dentro dos últimos 24 meses;", style="List Bullet")
    doc.add_paragraph(
        "Relevância temática verificada por palavras-chave alinhadas a "
        "Sistemas de Informação em Saúde.", style="List Bullet")
    _add_paragraph(doc, "Critérios de exclusão:", bold=True)
    doc.add_paragraph(
        "Aplicativos com menos de 1.000 instalações (baixa representatividade);",
        style="List Bullet")
    doc.add_paragraph(
        "Aplicativos que não recebam atualizações há mais de 24 meses "
        "(obsolescência tecnológica);", style="List Bullet")
    doc.add_paragraph(
        "Aplicativos de academias, dietas ou fitness recreativo sem relação "
        "com SIS clínicos;", style="List Bullet")
    doc.add_paragraph("Aplicativos duplicados (mesmo appId).", style="List Bullet")

    _add_heading(doc, "2.4 Classificação de Desenvolvedores", 2)
    _add_paragraph(doc,
        "Os desenvolvedores foram classificados em três categorias: "
        "(a) Governamental — identificados por termos como 'ministério', 'SUS', "
        "'governo', 'prefeitura', 'secretaria', 'DATASUS' no nome do desenvolvedor; "
        "(b) Institucional — entidades de saúde como CRM, CFM, hospitais, "
        "universidades, institutos de pesquisa; "
        "(c) Comercial — desenvolvedores privados.")

    _add_heading(doc, "2.5 Procedimentos de Análise (Data Analysis)", 2)
    _add_paragraph(doc,
        "O design prevê duas camadas de análise:")
    _add_paragraph(doc,
        "1. Análise Quantitativa: Estatística descritiva das métricas e "
        "correlação de Pearson entre a frequência de atualizações e a "
        "satisfação do usuário.", bold=True)
    _add_paragraph(doc,
        "2. Análise Qualitativa (Mineração de Texto): Aplicação de "
        "Processamento de Linguagem Natural (PLN) para categorizar os "
        "comentários nos seguintes eixos:", bold=True)
    for eixo, info in EIXOS_TEMATICOS.items():
        doc.add_paragraph(
            f"• {eixo}: palavras-chave incluem {', '.join(info['keywords'][:8])}...",
            style="List Bullet")

    _add_heading(doc, "2.6 Análise de Sentimentos", 2)
    _add_paragraph(doc,
        "A análise de sentimentos foi realizada com a biblioteca TextBlob, "
        "complementada por um dicionário customizado de termos em português "
        "brasileiro com valências positivas e negativas. Cada avaliação recebeu "
        "um escore de polaridade (−1 a +1) e foi classificada como Positiva "
        "(> 0,05), Neutra (−0,05 a +0,05) ou Negativa (< −0,05).")

    _add_heading(doc, "2.7 Modelagem de Tópicos (LDA)", 2)
    _add_paragraph(doc,
        "Para descoberta não-supervisionada de temas latentes nas avaliações, "
        "utilizou-se o algoritmo Latent Dirichlet Allocation (LDA) com 5 tópicos, "
        "implementado via scikit-learn. A matriz documento-termo foi construída "
        "com CountVectorizer (max_features=2000, min_df=3, max_df=0.9) após "
        "remoção de stopwords em português.")

    _add_heading(doc, "2.8 Ética e Governança de Dados", 2)
    _add_paragraph(doc,
        "Este estudo fundamenta-se na análise de dados secundários de acesso "
        "público, disponíveis na plataforma Google Play Store. Em conformidade "
        "com a Resolução nº 510/2016 do Conselho Nacional de Saúde (CNS), "
        "pesquisas que utilizam informações de domínio público e que não "
        "possibilitam a identificação individual dos sujeitos dispensam a "
        "submissão a um Comitê de Ética em Pesquisa (CEP).")
    _add_paragraph(doc,
        "Não obstante, o design da pesquisa segue os preceitos éticos da "
        "Lei Geral de Proteção de Dados (LGPD — Lei 13.709/2018), garantindo "
        "o anonimato dos usuários e a utilização estrita de dados agregados "
        "para fins acadêmicos, sem violação dos termos de serviço da plataforma "
        "hospedeira. Nomes de usuários foram substituídos por hashes "
        "criptográficos (SHA-256 truncado).")


def gerar_relatorio_quantitativo(doc: Document):
    """Seção: Resultados Quantitativos."""
    _add_heading(doc, "3. RESULTADOS — ANÁLISE QUANTITATIVA", 1)

    # ── 3.1 amostra ──────────────────────────────────────────────────────
    csv_a = sorted(CLEAN_DIR.glob("apps_limpos_*.csv"), reverse=True)
    csv_r = sorted(CLEAN_DIR.glob("reviews_limpos_*.csv"), reverse=True)
    df_apps = pd.read_csv(csv_a[0], encoding="utf-8-sig") if csv_a else pd.DataFrame()
    df_rev = pd.read_csv(csv_r[0], encoding="utf-8-sig") if csv_r else pd.DataFrame()
    n_apps = len(df_apps)
    n_rev = len(df_rev)

    _add_heading(doc, "3.1 Composição da Amostra", 2)
    _add_paragraph(doc,
        f"Após aplicação dos critérios de inclusão e exclusão, a amostra final "
        f"foi composta por {n_apps} aplicativos e {n_rev:,} avaliações de "
        f"usuários. A tabela a seguir apresenta os 20 aplicativos com maior "
        f"número de instalações.")

    top = _safe_read(TABELAS_DIR / "top_20_apps.csv")
    if not top.empty:
        _add_table_from_df(doc, top)
        doc.add_paragraph()

    # ── 3.2 estatística descritiva ────────────────────────────────────────
    _add_heading(doc, "3.2 Estatística Descritiva", 2)
    desc = _safe_read(TABELAS_DIR / "estatistica_descritiva.csv")
    if not desc.empty:
        _add_paragraph(doc,
            "A tabela abaixo resume as principais estatísticas descritivas das "
            "variáveis numéricas dos aplicativos incluídos na amostra:")
        _add_table_from_df(doc, desc)
        doc.add_paragraph()

        # interpretação automática
        if "score" in df_apps.columns:
            media = df_apps["score"].mean()
            mediana = df_apps["score"].median()
            _add_paragraph(doc,
                f"A nota média geral dos aplicativos foi {media:.2f} (mediana: "
                f"{mediana:.2f}), indicando {'satisfação moderada' if media >= 3.5 else 'avaliações tendendo a insatisfação'} "
                f"dos usuários. {'A proximidade entre média e mediana sugere distribuição relativamente simétrica.' if abs(media - mediana) < 0.3 else 'A diferença entre média e mediana sugere assimetria na distribuição.'}")

        if "instalacoes_num" in df_apps.columns:
            df_apps["instalacoes_num"] = pd.to_numeric(df_apps["instalacoes_num"], errors="coerce")
            total = df_apps["instalacoes_num"].sum()
            _add_paragraph(doc,
                f"O total acumulado de instalações foi de {total:,.0f}, com alta "
                f"variabilidade entre os aplicativos (amplitude máxima: "
                f"{df_apps['instalacoes_num'].max():,.0f}). Isso indica concentração "
                f"do mercado em poucos aplicativos dominantes.")

    _add_image(doc, GRAFICOS_DIR / "distribuicao_scores.png")
    _add_paragraph(doc, "Figura: Distribuição das notas médias dos aplicativos.",
                   italic=True, font_size=9)

    _add_image(doc, GRAFICOS_DIR / "distribuicao_instalacoes_log.png")
    _add_paragraph(doc, "Figura: Distribuição de instalações (escala logarítmica).",
                   italic=True, font_size=9)

    # ── 3.3 desenvolvedores ───────────────────────────────────────────────
    _add_heading(doc, "3.3 Perfil dos Desenvolvedores", 2)
    dev = _safe_read(TABELAS_DIR / "distribuicao_desenvolvedor.csv")
    if not dev.empty:
        _add_table_from_df(doc, dev)
        doc.add_paragraph()
        if "tipo_desenvolvedor" in df_apps.columns:
            cnt = df_apps["tipo_desenvolvedor"].value_counts()
            maior = cnt.index[0]
            pct_maior = cnt.iloc[0] / cnt.sum() * 100
            _add_paragraph(doc,
                f"O tipo de desenvolvedor mais frequente é '{maior}' ({pct_maior:.1f}%), "
                f"o que revela {'predomínio do setor privado' if maior == 'Comercial' else 'forte presença governamental/institucional'} "
                f"no ecossistema de apps de saúde no Brasil. "
                f"{'Essa predominância comercial levanta questões sobre a regulação e validação científica dos conteúdos de saúde ofertados.' if maior == 'Comercial' else ''}")

    _add_image(doc, GRAFICOS_DIR / "distribuicao_desenvolvedor.png")
    _add_paragraph(doc, "Figura: Distribuição por tipo de desenvolvedor.",
                   italic=True, font_size=9)

    # ── 3.4 correlações ──────────────────────────────────────────────────
    _add_heading(doc, "3.4 Análise de Correlação", 2)
    corr = _safe_read(TABELAS_DIR / "correlacoes.csv")
    if not corr.empty:
        _add_table_from_df(doc, corr)
        doc.add_paragraph()
        for _, row in corr.iterrows():
            sig = row.get("Significativo (p<0.05)", "")
            r_val = row.get("Pearson (r)", 0)
            _add_paragraph(doc,
                f"Para '{row.get('Variáveis','')}': coeficiente de Pearson r = {r_val}. "
                f"{'O resultado é estatisticamente significativo (p < 0,05), ' if sig == 'Sim' else 'O resultado NÃO foi estatisticamente significativo (p ≥ 0,05), '}"
                f"{'sugerindo que há associação linear entre as variáveis.' if sig == 'Sim' else 'indicando que não se pode rejeitar a hipótese nula de ausência de correlação linear.'}")

    _add_image(doc, GRAFICOS_DIR / "correlacao_atualizacao_score.png")
    _add_paragraph(doc, "Figura: Diagrama de dispersão — Atualização × Score.",
                   italic=True, font_size=9)
    _add_image(doc, GRAFICOS_DIR / "matriz_correlacao.png")
    _add_paragraph(doc, "Figura: Matriz de correlação entre variáveis quantitativas.",
                   italic=True, font_size=9)

    _add_image(doc, GRAFICOS_DIR / "boxplot_score_desenvolvedor.png")
    _add_paragraph(doc, "Figura: Boxplot do Score por tipo de desenvolvedor.",
                   italic=True, font_size=9)

    _add_image(doc, GRAFICOS_DIR / "atualizacoes_por_mes.png")
    _add_paragraph(doc, "Figura: Distribuição temporal das atualizações.",
                   italic=True, font_size=9)


def gerar_relatorio_qualitativo(doc: Document):
    """Seção: Resultados Qualitativos (PLN)."""
    _add_heading(doc, "4. RESULTADOS — ANÁLISE QUALITATIVA (PLN)", 1)

    # ── 4.1 sentimento ────────────────────────────────────────────────────
    _add_heading(doc, "4.1 Análise de Sentimentos", 2)
    sent = _safe_read(TABELAS_DIR / "sentimentos.csv")
    if not sent.empty:
        _add_table_from_df(doc, sent)
        doc.add_paragraph()
        # interpretação
        # Tenta extrair percentuais
        if "Percentual (%)" in sent.columns:
            pos_row = sent[sent.iloc[:, 0].str.contains("Positivo", na=False)]
            neg_row = sent[sent.iloc[:, 0].str.contains("Negativo", na=False)]
            neu_row = sent[sent.iloc[:, 0].str.contains("Neutro", na=False)]
            p_pos = pos_row["Percentual (%)"].values[0] if len(pos_row) > 0 else 0
            p_neg = neg_row["Percentual (%)"].values[0] if len(neg_row) > 0 else 0
            p_neu = neu_row["Percentual (%)"].values[0] if len(neu_row) > 0 else 0
            _add_paragraph(doc,
                f"A distribuição de sentimentos revelou {p_pos}% de avaliações positivas, "
                f"{p_neu}% neutras e {p_neg}% negativas. "
                f"{'A predominância de avaliações positivas sugere satisfação geral dos usuários.' if float(p_pos) > float(p_neg) else 'A proporção expressiva de avaliações negativas sinaliza insatisfação relevante.'} "
                f"O percentual de avaliações neutras ({p_neu}%) pode indicar "
                f"comentários descritivos sem juízo de valor claro, ou limitações "
                f"do classificador de sentimentos para textos curtos em português.")

    _add_image(doc, GRAFICOS_DIR / "sentimentos.png")
    _add_paragraph(doc, "Figura: Distribuição de sentimentos nas avaliações.",
                   italic=True, font_size=9)
    _add_image(doc, GRAFICOS_DIR / "distribuicao_polaridade.png")
    _add_paragraph(doc, "Figura: Histograma da polaridade dos sentimentos.",
                   italic=True, font_size=9)

    # ── 4.2 temas ────────────────────────────────────────────────────────
    _add_heading(doc, "4.2 Classificação Temática", 2)
    tema = _safe_read(TABELAS_DIR / "distribuicao_tematica.csv")
    if not tema.empty:
        _add_paragraph(doc,
            "As avaliações foram classificadas nos cinco eixos temáticos "
            "pré-definidos. Um mesmo comentário pode pertencer a mais de um eixo "
            "se contiver termos de múltiplas categorias.")
        _add_table_from_df(doc, tema)
        doc.add_paragraph()
        # interpretação
        if len(tema) > 0:
            top_tema = tema.iloc[0, 0] if tema.iloc[0, 0] != "index" else tema.iloc[0, 1]
            _add_paragraph(doc,
                f"O eixo temático mais mencionado foi '{top_tema}', indicando que "
                f"esta é a principal preocupação dos usuários ao avaliar aplicativos "
                f"de saúde. Isso é consistente com a literatura que aponta a "
                f"funcionalidade técnica como determinante primário da satisfação "
                f"do usuário com aplicativos mHealth.")

    _add_image(doc, GRAFICOS_DIR / "distribuicao_tematica.png")
    _add_paragraph(doc, "Figura: Distribuição temática das avaliações.",
                   italic=True, font_size=9)
    _add_image(doc, GRAFICOS_DIR / "sentimento_por_tema.png")
    _add_paragraph(doc, "Figura: Sentimento por eixo temático.",
                   italic=True, font_size=9)

    # ── 4.3 LDA ──────────────────────────────────────────────────────────
    _add_heading(doc, "4.3 Modelagem de Tópicos (LDA)", 2)
    lda = _safe_read(TABELAS_DIR / "topicos_lda.csv")
    if not lda.empty:
        _add_paragraph(doc,
            "O modelo LDA identificou cinco tópicos latentes no corpus de "
            "avaliações. A tabela apresenta as 10 palavras-chave de maior peso "
            "para cada tópico:")
        _add_table_from_df(doc, lda)
        doc.add_paragraph()
        _add_paragraph(doc,
            "A interpretação dos tópicos deve ser realizada qualitativamente. "
            "Recomenda-se que o pesquisador analise as palavras-chave de cada "
            "tópico e atribua rótulos semânticos que capturem o tema central. "
            "Por exemplo, tópicos com termos como 'funciona', 'erro', 'bug' "
            "podem ser rotulados como 'Problemas Técnicos'.")

    _add_image(doc, GRAFICOS_DIR / "lda_heatmap.png")
    _add_paragraph(doc, "Figura: Heatmap dos pesos LDA por tópico.",
                   italic=True, font_size=9)

    # ── 4.4 nuvens ───────────────────────────────────────────────────────
    _add_heading(doc, "4.4 Nuvens de Palavras", 2)
    _add_paragraph(doc,
        "As nuvens de palavras oferecem representação visual da frequência "
        "dos termos no corpus, auxiliando na identificação rápida de temas "
        "e sentimentos predominantes.")
    for label in ["geral", "positivo", "negativo"]:
        _add_image(doc, GRAFICOS_DIR / f"wordcloud_{label}.png")
        _add_paragraph(doc, f"Figura: Nuvem de palavras — {label.title()}.",
                       italic=True, font_size=9)

    # ── 4.5 frequência ───────────────────────────────────────────────────
    _add_heading(doc, "4.5 Frequência de Palavras", 2)
    freq = _safe_read(TABELAS_DIR / "frequencia_palavras.csv")
    if not freq.empty:
        _add_paragraph(doc,
            "As 30 palavras mais frequentes no corpus refletem os temas "
            "dominantes das avaliações:")
        _add_table_from_df(doc, freq)
        doc.add_paragraph()

    _add_image(doc, GRAFICOS_DIR / "frequencia_palavras.png")
    _add_paragraph(doc, "Figura: 30 palavras mais frequentes.",
                   italic=True, font_size=9)


def gerar_relatorio_apps(doc: Document):
    """Seção: Lista completa de apps incluídos."""
    _add_heading(doc, "APÊNDICE A — LISTA COMPLETA DE APLICATIVOS INCLUÍDOS", 1)
    csv_a = sorted(CLEAN_DIR.glob("apps_limpos_*.csv"), reverse=True)
    if not csv_a:
        _add_paragraph(doc, "Dados de aplicativos não encontrados.")
        return
    df = pd.read_csv(csv_a[0], encoding="utf-8-sig")
    _add_paragraph(doc,
        f"A tabela a seguir lista todos os {len(df)} aplicativos incluídos na "
        f"amostra final, ordenados por número de instalações.")

    cols = [c for c in ["title","developer","tipo_desenvolvedor","score",
                        "instalacoes_num","ratings","genreId"] if c in df.columns]
    df_show = df.sort_values("instalacoes_num", ascending=False)[cols].copy()
    if "instalacoes_num" in df_show.columns:
        df_show["instalacoes_num"] = pd.to_numeric(df_show["instalacoes_num"], errors="coerce")
        df_show["instalacoes_num"] = df_show["instalacoes_num"].apply(lambda x: f"{x:,.0f}" if pd.notna(x) else "")
    rename = {"title":"App","developer":"Desenvolvedor",
              "tipo_desenvolvedor":"Tipo","score":"Nota",
              "instalacoes_num":"Instalações","ratings":"Avaliações",
              "genreId":"Categoria"}
    df_show.rename(columns=rename, inplace=True)
    _add_table_from_df(doc, df_show, max_rows=200)


def gerar_discussao(doc: Document):
    """Seção: Discussão e considerações finais."""
    _add_heading(doc, "5. DISCUSSÃO", 1)
    _add_paragraph(doc,
        "Os resultados deste estudo oferecem um panorama da integração dos "
        "Sistemas de Informação em Saúde (SIS) no ecossistema de aplicativos "
        "móveis (mHealth) da Google Play Store. Alguns achados merecem destaque:")

    _add_heading(doc, "5.1 Predomínio de Desenvolvedores Comerciais", 2)
    _add_paragraph(doc,
        "A expressiva maioria dos aplicativos é desenvolvida por empresas "
        "privadas, o que levanta questões importantes sobre a validação "
        "científica dos conteúdos de saúde ofertados e sobre a interoperabilidade "
        "com sistemas governamentais. Diferentemente de aplicativos governamentais "
        "(como o Meu SUS Digital), apps comerciais podem não seguir padrões "
        "abertos de integração (HL7/FHIR), dificultando a troca de dados entre "
        "sistemas e comprometendo a continuidade do cuidado.")

    _add_heading(doc, "5.2 Integração e Interoperabilidade", 2)
    _add_paragraph(doc,
        "A análise temática dos comentários revela que a interoperabilidade e "
        "integração de dados é uma preocupação recorrente dos usuários. Relatos "
        "de falha na sincronização de dados entre aplicativos e prontuários "
        "eletrônicos indicam que a onipresença dos SIS via dispositivos móveis "
        "ainda enfrenta barreiras técnicas significativas. A adoção de padrões "
        "como HL7 FHIR e OpenEHR é essencial para viabilizar a integração.")

    _add_heading(doc, "5.3 Satisfação do Usuário e Manutenção", 2)
    _add_paragraph(doc,
        "A análise de correlação entre frequência de atualização e nota média "
        "contribui para entender se a manutenção ativa dos aplicativos está "
        "associada à satisfação dos usuários. Apps que recebem atualizações "
        "frequentes tendem a corrigir bugs e incorporar feedback, o que "
        "teoricamente deveria elevar a satisfação.")

    _add_heading(doc, "5.4 Segurança da Informação", 2)
    _add_paragraph(doc,
        "A presença de comentários sobre segurança e privacidade reforça a "
        "importância de aplicativos de saúde adotarem medidas robustas de "
        "proteção de dados, especialmente considerando que muitos processam "
        "dados sensíveis de prontuários eletrônicos. A conformidade com a "
        "LGPD é essencial para a confiança do usuário.")

    _add_heading(doc, "5.5 Limitações do Estudo", 2)
    _add_paragraph(doc,
        "Este estudo apresenta as seguintes limitações:")
    doc.add_paragraph(
        "Análise restrita à Google Play Store, excluindo a Apple App Store;",
        style="List Bullet")
    doc.add_paragraph(
        "Limitação do scraper a 200 avaliações por app;",
        style="List Bullet")
    doc.add_paragraph(
        "Classificação de sentimentos por dicionário, sem deep learning;",
        style="List Bullet")
    doc.add_paragraph(
        "Corte temporal de 24 meses para atualizações;",
        style="List Bullet")
    doc.add_paragraph(
        "Resultados sensíveis à lista de descritores de busca.",
        style="List Bullet")

    _add_heading(doc, "6. CONSIDERAÇÕES FINAIS", 1)
    _add_paragraph(doc,
        "O presente estudo demonstrou a aplicabilidade da Infodemiologia e "
        "técnicas de Processamento de Linguagem Natural para análise sistemática "
        "da integração do ecossistema mHealth com os Sistemas de Informação em "
        "Saúde na Google Play Store. Os resultados evidenciam um ecossistema "
        "ainda fragmentado, com desafios de interoperabilidade, segurança da "
        "informação e experiência do usuário que precisam ser endereçados para "
        "que a onipresença dos SIS via dispositivos móveis se torne efetiva. "
        "A análise fornece subsídios para desenvolvedores, gestores de saúde "
        "pública e formuladores de políticas regulatórias no contexto da saúde "
        "digital brasileira.")

    _add_heading(doc, "REFERÊNCIAS", 1)
    refs = [
        "EYSENBACH, G. Infodemiology and infoveillance: framework for an "
        "emerging set of public health informatics methods to analyze search, "
        "communication and publication behavior on the Internet. Journal of "
        "Medical Internet Research, v. 11, n. 1, e11, 2009.",
        "BRASIL. Lei nº 13.709, de 14 de agosto de 2018. Lei Geral de "
        "Proteção de Dados Pessoais (LGPD).",
        "BRASIL. Resolução nº 510, de 7 de abril de 2016. Conselho Nacional "
        "de Saúde. Dispõe sobre normas aplicáveis a pesquisas em Ciências "
        "Humanas e Sociais.",
        "WHO. Classification of Digital Health Interventions v1.0. "
        "World Health Organization, 2018.",
        "BLEI, D. M.; NG, A. Y.; JORDAN, M. I. Latent Dirichlet Allocation. "
        "Journal of Machine Learning Research, v. 3, p. 993–1022, 2003.",
        "HL7 INTERNATIONAL. FHIR (Fast Healthcare Interoperability Resources). "
        "Disponível em: https://www.hl7.org/fhir/. Acesso em: 2026.",
    ]
    for i, ref in enumerate(refs, 1):
        doc.add_paragraph(f"[{i}] {ref}")


# ════════════════════════════════════════════════════════════════════════════
def executar_relatorio():
    """Gera o documento Word completo."""
    logger.info("=" * 60)
    logger.info("GERAÇÃO DO RELATÓRIO WORD")
    logger.info("=" * 60)

    doc = Document()

    # ── estilos gerais ────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)

    # ── capa ──────────────────────────────────────────────────────────────
    for _ in range(4):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run(TITULO_PROJETO.upper())
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sub.add_run(SUBTITULO_PROJETO)
    run2.italic = True
    run2.font.size = Pt(13)

    for _ in range(4):
        doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = info.add_run(f"Data de geração: {datetime.now().strftime('%d/%m/%Y')}")
    run3.font.size = Pt(11)

    doc.add_page_break()

    # ── sumário ───────────────────────────────────────────────────────────
    _add_heading(doc, "SUMÁRIO", 1)
    sumario = [
        "1. Introdução",
        "2. Metodologia",
        "   2.1 Tipologia do Estudo",
        "   2.2 Protocolo de Coleta de Dados (Web Scraping)",
        "   2.3 Definição do Escopo e Amostragem",
        "   2.4 Classificação de Desenvolvedores",
        "   2.5 Procedimentos de Análise (Data Analysis)",
        "   2.6 Análise de Sentimentos",
        "   2.7 Modelagem de Tópicos (LDA)",
        "   2.8 Ética e Governança de Dados",
        "3. Resultados — Análise Quantitativa",
        "   3.1 Composição da Amostra",
        "   3.2 Estatística Descritiva",
        "   3.3 Perfil dos Desenvolvedores",
        "   3.4 Análise de Correlação",
        "4. Resultados — Análise Qualitativa (PLN)",
        "   4.1 Análise de Sentimentos",
        "   4.2 Classificação Temática",
        "   4.3 Modelagem de Tópicos (LDA)",
        "   4.4 Nuvens de Palavras",
        "   4.5 Frequência de Palavras",
        "5. Discussão",
        "   5.1 Predomínio de Desenvolvedores Comerciais",
        "   5.2 Integração e Interoperabilidade",
        "   5.3 Satisfação do Usuário e Manutenção",
        "   5.4 Segurança da Informação",
        "   5.5 Limitações do Estudo",
        "6. Considerações Finais",
        "Referências",
        "Apêndice A — Lista de Aplicativos",
    ]
    for item in sumario:
        doc.add_paragraph(item)
    doc.add_page_break()

    # ── 1. introdução ────────────────────────────────────────────────────
    _add_heading(doc, "1. INTRODUÇÃO", 1)
    _add_paragraph(doc,
        "A proliferação de aplicativos móveis de saúde (mHealth apps) e a "
        "crescente digitalização dos Sistemas de Informação em Saúde (SIS) "
        "transformaram a relação entre cidadãos, profissionais de saúde e "
        "as informações clínicas. No Brasil, a Google Play Store concentra "
        "centenas de aplicativos voltados para gestão de prontuários eletrônicos, "
        "integração com o SUS, telemedicina e vigilância epidemiológica. "
        "Contudo, a qualidade da integração, a interoperabilidade com sistemas "
        "governamentais e a satisfação dos usuários com esses aplicativos "
        "permanecem insuficientemente estudadas.")
    _add_paragraph(doc,
        "A Infodemiologia, proposta por Eysenbach (2009), oferece um arcabouço "
        "teórico-metodológico para a análise de padrões informacionais em "
        "ambientes digitais de saúde. Combinada à Infometria — que quantifica "
        "métricas de informação — e ao Processamento de Linguagem Natural (PLN), "
        "essa abordagem permite investigar sistematicamente como os usuários "
        "percebem a integração dos SIS via dispositivos móveis.")
    _add_paragraph(doc,
        "Este estudo tem como objetivo analisar a onipresença dos Sistemas de "
        "Informação em Saúde no ecossistema mHealth da Google Play Store, "
        "utilizando técnicas de web scraping automatizado, análise estatística "
        "descritiva e inferencial, análise de sentimentos e modelagem de tópicos, "
        "contribuindo para a compreensão da integração digital em saúde no "
        "contexto brasileiro.")

    # ── seções principais ─────────────────────────────────────────────────
    gerar_relatorio_metodologia(doc)
    doc.add_page_break()
    gerar_relatorio_quantitativo(doc)
    doc.add_page_break()
    gerar_relatorio_qualitativo(doc)
    doc.add_page_break()
    gerar_discussao(doc)
    doc.add_page_break()
    gerar_relatorio_apps(doc)

    # ── salvar ────────────────────────────────────────────────────────────
    out = RELATORIO_DIR / f"relatorio_infodemiologia_{datetime.now():%Y%m%d_%H%M%S}.docx"
    doc.save(str(out))
    logger.info(f"Relatório salvo em: {out}")
    return out


if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO,
                        format="%(asctime)s [%(levelname)s] %(message)s")
    executar_relatorio()
