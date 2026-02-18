# -*- coding: utf-8 -*-
"""
Fase 2.5 — Geração de lista Word de todos os apps coletados para revisão.

Após a coleta (Fases 1+2), gera um documento Word com TODOS os aplicativos
encontrados na Google Play Store, incluindo metadados principais e uma coluna
"Incluir?" em branco para o pesquisador anotar manualmente quais entrarão
na análise.

O arquivo gerado fica em:
  resultados/relatorios/lista_apps_para_revisao_YYYYMMDD_HHMMSS.docx

Uso isolado:
  python -m src.lista_coleta
"""

import logging
from datetime import datetime
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from src.config import RAW_DIR, RELATORIO_DIR

logger = logging.getLogger(__name__)

# ── Colunas a exibir ─────────────────────────────────────────────────────────
COLUNAS = [
    ("Nº",              None,            0.6),   # índice sequencial
    ("App ID",          "appId",         3.2),
    ("Nome",            "title",         3.8),
    ("Desenvolvedor",   "developer",     3.5),
    ("Categoria",       "genreId",       2.2),
    ("Nota",            "score",         1.0),
    ("Instalações",     "installs",      2.0),
    ("Atualizado em",   "lastUpdatedOn", 2.5),
    ("Gratuito",        "free",          1.2),
    ("Incluir?",        None,            1.2),   # coluna em branco para marcação
]

# larguras em cm
_WIDTHS_CM = [c[2] for c in COLUNAS]


# ── helpers de formatação ─────────────────────────────────────────────────────
def _shade_cell(cell, fill: str):
    tcp = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcp.append(shd)


def _set_cell_text(cell, text: str, bold=False, size=8,
                    color=None, align=WD_ALIGN_PARAGRAPH.LEFT,
                    fill: str | None = None):
    if fill:
        _shade_cell(cell, fill)
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)


def _set_col_width(cell, width_cm: float):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(int(width_cm * 567)))   # 1cm ≈ 567 twips
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


# ── leitura dos dados brutos ──────────────────────────────────────────────────
def _carregar_brutos() -> pd.DataFrame:
    csvs = sorted(RAW_DIR.glob("apps_brutos_*.csv"), reverse=True)
    if not csvs:
        raise FileNotFoundError(f"Nenhum arquivo apps_brutos_*.csv em {RAW_DIR}")
    df = pd.read_csv(csvs[0], encoding="utf-8-sig")
    logger.info(f"Apps brutos carregados: {len(df)} (de {csvs[0].name})")
    return df


# ── geração do documento Word ─────────────────────────────────────────────────
def gerar_lista_apps_word(df: pd.DataFrame | None = None) -> Path:
    """
    Gera o Word de revisão.

    Parâmetros
    ----------
    df : DataFrame opcional. Se None, carrega o CSV de apps_brutos mais recente.

    Retorna
    -------
    Path do arquivo gerado.
    """
    if df is None:
        df = _carregar_brutos()

    # ordenar por nome
    if "title" in df.columns:
        df = df.sort_values("title", ignore_index=True)

    total = len(df)
    doc = Document()

    # margens
    for sec in doc.sections:
        sec.orientation = 1                    # landscape
        sec.page_width, sec.page_height = sec.page_height, sec.page_width
        sec.top_margin    = Cm(1.5)
        sec.bottom_margin = Cm(1.5)
        sec.left_margin   = Cm(1.8)
        sec.right_margin  = Cm(1.8)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(9)

    # ── capa ─────────────────────────────────────────────────────────────────
    for _ in range(3):
        doc.add_paragraph()

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("LISTA DE APLICATIVOS COLETADOS\nPARA REVISÃO E SELEÇÃO")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run(
        "Sistemas de Informação em Saúde — Integração mHealth via Google Play Store\n"
        "Marque a coluna 'Incluir?' para os apps que entrarão na análise.")
    r2.italic = True
    r2.font.size = Pt(12)

    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = info.add_run(
        f"Data de geração: {datetime.now().strftime('%d/%m/%Y %H:%M')}     "
        f"Total de aplicativos: {total}")
    r3.font.size = Pt(10)

    doc.add_page_break()

    # ── instrução ─────────────────────────────────────────────────────────────
    inst = doc.add_paragraph()
    ri = inst.add_run(
        "INSTRUÇÕES:  Analise os aplicativos abaixo e marque '✓' na coluna "
        "'Incluir?' para aqueles que devem entrar na análise. "
        "A coluna 'Nota' vai de 0 a 5. "
        "Após concluir, use o número (Nº) do app no pipeline de seleção.")
    ri.italic = True
    ri.font.size = Pt(9)
    ri.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    inst.paragraph_format.space_after = Pt(8)

    # ── tabela ────────────────────────────────────────────────────────────────
    n_cols = len(COLUNAS)
    table = doc.add_table(rows=1, cols=n_cols, style="Table Grid")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # cabeçalho
    hdr = table.rows[0].cells
    for j, (label, _, w) in enumerate(COLUNAS):
        _set_col_width(hdr[j], w)
        _set_cell_text(hdr[j], label,
                       bold=True, size=8,
                       color="FFFFFF",
                       align=WD_ALIGN_PARAGRAPH.CENTER,
                       fill="1A237E")

    # linhas de dados
    for i, (_, row) in enumerate(df.iterrows(), start=1):
        cells = table.add_row().cells
        fill  = "EEF2FF" if i % 2 == 0 else "FFFFFF"

        def val(field, default=""):
            v = row.get(field, default)
            return "" if pd.isna(v) else str(v)

        # Nº
        _set_cell_text(cells[0], str(i),
                       align=WD_ALIGN_PARAGRAPH.CENTER, fill=fill)
        # App ID
        txt_id = val("appId")
        _set_cell_text(cells[1], txt_id[:50], fill=fill)
        # Nome
        _set_cell_text(cells[2], val("title")[:55], fill=fill)
        # Desenvolvedor
        _set_cell_text(cells[3], val("developer")[:50], fill=fill)
        # Categoria
        _set_cell_text(cells[4], val("genreId"), fill=fill)
        # Nota
        score_raw = row.get("score")
        if pd.notna(score_raw):
            try:
                score_str = f"{float(score_raw):.1f}"
            except (ValueError, TypeError):
                score_str = val("score")
        else:
            score_str = "—"
        _set_cell_text(cells[5], score_str,
                       align=WD_ALIGN_PARAGRAPH.CENTER, fill=fill)
        # Instalações
        _set_cell_text(cells[6], val("installs") or val("minInstalls"), fill=fill)
        # Atualizado em
        _set_cell_text(cells[7],
                       val("lastUpdatedOn") or val("updated"), fill=fill)
        # Gratuito
        free_raw = row.get("free")
        free_str = "Sim" if str(free_raw).lower() in ("true","1","sim","yes") else (
                   "Não" if pd.notna(free_raw) else "—")
        _set_cell_text(cells[8], free_str,
                       align=WD_ALIGN_PARAGRAPH.CENTER, fill=fill)
        # Incluir? (em branco, fundo levemente diferente para destacar)
        _set_cell_text(cells[9], "",
                       fill="FFF9C4")   # amarelo claro = campo a preencher

    # definir larguras de todas as células
    for row_obj in table.rows:
        for j, (_, _, w) in enumerate(COLUNAS):
            _set_col_width(row_obj.cells[j], w)

    # ── rodapé / instrução extra ──────────────────────────────────────────────
    doc.add_paragraph()
    rod = doc.add_paragraph()
    rod.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = rod.add_run(
        f"Total: {total} aplicativos coletados   |   "
        "Pipeline — Sistemas de Informação em Saúde / mHealth   |   "
        f"{datetime.now().strftime('%d/%m/%Y')}")
    rr.font.size = Pt(8)
    rr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    rr.italic = True

    # ── salvar ────────────────────────────────────────────────────────────────
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    out = RELATORIO_DIR / f"lista_apps_para_revisao_{ts}.docx"
    doc.save(str(out))
    logger.info(f"Lista de revisão gerada: {out}")
    print(f"\n  ✔ Lista Word gerada → {out}")
    return out


def executar_lista_coleta(df: pd.DataFrame | None = None) -> Path:
    """Ponto de entrada chamado pelo pipeline."""
    logger.info("=" * 60)
    logger.info("FASE 2.5 — Lista Word de Apps para Revisão")
    logger.info("=" * 60)
    return gerar_lista_apps_word(df)


if __name__ == "__main__":
    import sys
    logging.basicConfig(level=logging.INFO,
                        format="%(asctime)s [%(levelname)s] %(message)s")
    sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
    executar_lista_coleta()
